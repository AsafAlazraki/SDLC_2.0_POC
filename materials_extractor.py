"""
materials_extractor.py — turn any uploaded material into the text the agents see.

Strategy: dispatch by file extension / MIME, return (extracted_text, metadata).
Never raises — extraction failures yield ("", {"error": ...}) so the upload
still records and the user gets an honest record of what couldn't be parsed.

Special handling:
  - .oap files (OutSystems Application Packages) are zip archives. We crack the
    archive open, read Manifest.xml, list the .oml modules (binary, can't be
    parsed but their names + sizes are valuable signal), and concatenate any
    text-shaped contents up to a reasonable cap.
  - .zip / .gz tarballs: same recursion strategy as OAP.
  - .pdf via pypdf  (soft-imported; missing dep = clean error message).
  - .docx via python-docx (same).
  - Code / text / markup: decoded as UTF-8 with replacement.
  - Images: Gemini 2.0 Flash vision generates a structured text summary when
    GEMINI_API_KEY is set. Without a key, we fall back to metadata-only.
  - Audio (mp3/wav/m4a/ogg/flac/webm): Gemini 2.0 Flash transcribes and
    summarises natively when GEMINI_API_KEY is set. Same graceful fallback.

Hard cap: MAX_TEXT_BYTES per extracted blob to protect prompt budget.
"""

from __future__ import annotations
import base64
import io
import logging
import os
import zipfile
from typing import Tuple, Dict, Any, List

logger = logging.getLogger(__name__)

# Optional libraries — soft-imported so missing deps fail gracefully per file.
try:
    import pypdf  # type: ignore
except ImportError:
    pypdf = None  # type: ignore

try:
    import docx as docx_lib  # type: ignore  # python-docx
except ImportError:
    docx_lib = None  # type: ignore

# Gemini SDK is already a hard dep of the engine; treat it as optional here so
# materials_extractor stays usable in test contexts that mock it out.
try:
    from google import genai as _genai  # type: ignore
    from google.genai import types as _genai_types  # type: ignore
except ImportError:
    _genai = None  # type: ignore
    _genai_types = None  # type: ignore


MAX_TEXT_BYTES = 200_000        # per-material body cap (~50K tokens)
MAX_ZIP_FILES = 60              # cap how many entries we crack open inside a zip
MAX_ZIP_BYTES_PER_ENTRY = 80_000

# Vision extraction caps. Gemini 2.0 Flash accepts up to ~7MB per image but
# real screenshots/wireframes are typically <2MB. We cap at 4MB to stop a
# single oversized PNG from chewing the entire upload budget.
MAX_IMAGE_BYTES_FOR_VISION = 4 * 1024 * 1024
VISION_MODEL = "gemini-2.0-flash"

# Mapping of file extension → MIME type the Gemini SDK expects. SVGs are XML
# and we let the text path handle them — Gemini-vision rejects them anyway.
IMAGE_MIME_MAP = {
    ".png":  "image/png",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif":  "image/gif",
    ".webp": "image/webp",
    ".bmp":  "image/bmp",
}

# The vision prompt is intentionally structured. We want consistent fields the
# agents can rely on across mixed inputs (wireframes, screenshots, diagrams,
# whiteboard photos, ER diagrams, log captures...).
_VISION_PROMPT = """You are a visual analyst preparing this image for a fleet of software analysis agents who CANNOT see images. Your job is to convert it into structured text they can reason against.

Produce the following sections, omitting any that genuinely don't apply:

### What this image is
One short sentence: type of artefact (UI mockup, architecture diagram, ER diagram, screenshot of running app, whiteboard photo, code screenshot, log capture, error dialog, photograph, etc.).

### Visible content
A faithful, exhaustive description of what is in the image. For UI screenshots/wireframes: list every visible label, button, field, menu item, and their layout. For diagrams: list every node and every labelled edge/arrow. For text-bearing images (logs, code, error dialogs): transcribe the text verbatim where readable.

### Inferred purpose / context
What this artefact is trying to communicate, and the kind of system or workflow it implies.

### Notable details
Anything that stands out: error states, badges, version numbers, dates, brand marks, unusual layouts, accessibility concerns, anomalies.

### Open questions
2-3 things an analyst would need to ask the human to fully understand this image.

Be concise but complete. Do NOT speculate beyond what is visible. If the image is unreadable, blurry, or empty, say so plainly."""


async def extract_image_with_vision(
    payload: bytes,
    mime: str,
    *,
    gemini_api_key: str,
    filename: str = "",
) -> Tuple[str, Dict[str, Any]]:
    """Run Gemini-vision on an image and return (markdown_summary, metadata).

    Pure-async, never raises. Returns ("", {"error": ...}) on any failure so
    the caller can attach the metadata-only fallback path.

    Caller responsibility: only invoke when GEMINI_API_KEY is available and
    the image is under MAX_IMAGE_BYTES_FOR_VISION.
    """
    meta: Dict[str, Any] = {"vision_model": VISION_MODEL, "image_bytes": len(payload)}

    if not gemini_api_key:
        meta["error"] = "no Gemini API key — image kept as metadata only"
        return ("", meta)
    if _genai is None or _genai_types is None:
        meta["error"] = "google-genai SDK not installed"
        return ("", meta)
    if len(payload) > MAX_IMAGE_BYTES_FOR_VISION:
        meta["error"] = f"image too large ({len(payload):,} bytes > {MAX_IMAGE_BYTES_FOR_VISION:,})"
        return ("", meta)

    try:
        client = _genai.Client(api_key=gemini_api_key)
        image_part = _genai_types.Part.from_bytes(data=payload, mime_type=mime)
        response = await client.aio.models.generate_content(
            model=VISION_MODEL,
            contents=[image_part, _VISION_PROMPT],
            config=_genai_types.GenerateContentConfig(temperature=0.2),
        )
        summary = (response.text or "").strip()
        if not summary:
            meta["error"] = "vision returned empty response"
            return ("", meta)
        meta["extracted"] = "image_vision"
        meta["summary_chars"] = len(summary)
        # Wrap the summary so the agent prompt makes the provenance obvious.
        header = f"[Vision summary of image: {filename or '(unnamed)'} — {mime}]"
        return (f"{header}\n\n{summary}", meta)
    except Exception as e:
        logger.warning(f"Vision extraction failed for {filename}: {e}")
        meta["error"] = f"vision call failed: {e}"
        return ("", meta)


def _guess_image_mime(filename: str, mime: str) -> str:
    """Return the cleanest MIME for an image — prefer the supplied one, else
    derive from the extension. Defaults to image/png as Gemini's safest bet."""
    if mime and mime.lower().startswith("image/") and mime.lower() != "image/svg+xml":
        return mime.lower()
    ext = os.path.splitext(filename or "")[1].lower()
    return IMAGE_MIME_MAP.get(ext, "image/png")


# ─── Audio extraction (Phase 11) — symmetric to image vision ────────────────
# Gemini 2.0 Flash supports audio inputs natively via Part.from_bytes(mime).
# We send the audio with a structured prompt asking for a transcript +
# summary, just like the image vision path. Cap at 16MB (Gemini docs allow
# up to ~20MB but ~16 keeps single-file uploads safe with metadata overhead).

MAX_AUDIO_BYTES_FOR_TRANSCRIPTION = 16 * 1024 * 1024
AUDIO_MODEL = "gemini-2.0-flash"

_AUDIO_PROMPT = """You are a meeting transcription analyst preparing this audio for a fleet of software analysis agents who CANNOT listen to it. Your job is to convert it into structured text they can reason against.

Produce the following sections, omitting any that genuinely don't apply:

### What this audio is
One short sentence: type of recording (meeting, interview, voice memo, support call, demo walkthrough, podcast, lecture, other).

### Transcript
A faithful transcript. Identify speakers as "Speaker 1", "Speaker 2" etc. unless names are explicit. Use timestamps every 2-3 minutes if the audio is longer than 5 minutes. Quote verbatim where possible. Mark unintelligible passages as [unintelligible].

### Key topics discussed
A bulleted list of 3-8 main topics or decisions surfaced in the audio.

### Action items / commitments
Anyone who said "I will…" or who was assigned a task. Format: "Speaker N committed to X."

### Open questions / unresolved
Anything raised that did not get a clear answer or decision.

### Notable quotes
2-4 quotes that capture key positions, concerns, or decisions verbatim, with attribution.

Be concise but complete. Do NOT speculate beyond what is audible. If the audio is silent, garbled, or empty, say so plainly."""


async def extract_audio_with_gemini(
    payload: bytes,
    mime: str,
    *,
    gemini_api_key: str,
    filename: str = "",
) -> Tuple[str, Dict[str, Any]]:
    """Run Gemini audio understanding on an audio file. Returns
    (markdown_summary, metadata). Pure async, never raises.

    Mirrors `extract_image_with_vision()` — same graceful-degradation
    contract, same provenance header, same metadata-only fallback when no
    Gemini key is available.
    """
    meta: Dict[str, Any] = {"audio_model": AUDIO_MODEL, "audio_bytes": len(payload)}

    if not gemini_api_key:
        meta["error"] = "no Gemini API key — audio kept as metadata only"
        return ("", meta)
    if _genai is None or _genai_types is None:
        meta["error"] = "google-genai SDK not installed"
        return ("", meta)
    if len(payload) > MAX_AUDIO_BYTES_FOR_TRANSCRIPTION:
        meta["error"] = f"audio too large ({len(payload):,} bytes > {MAX_AUDIO_BYTES_FOR_TRANSCRIPTION:,})"
        return ("", meta)

    try:
        client = _genai.Client(api_key=gemini_api_key)
        audio_part = _genai_types.Part.from_bytes(data=payload, mime_type=mime)
        response = await client.aio.models.generate_content(
            model=AUDIO_MODEL,
            contents=[audio_part, _AUDIO_PROMPT],
            config=_genai_types.GenerateContentConfig(temperature=0.2),
        )
        summary = (response.text or "").strip()
        if not summary:
            meta["error"] = "audio transcription returned empty response"
            return ("", meta)
        meta["extracted"] = "audio_transcription"
        meta["summary_chars"] = len(summary)
        header = f"[Audio transcript + summary: {filename or '(unnamed)'} — {mime}]"
        return (f"{header}\n\n{summary}", meta)
    except Exception as e:
        logger.warning(f"Audio extraction failed for {filename}: {e}")
        meta["error"] = f"audio call failed: {e}"
        return ("", meta)


def _guess_audio_mime(filename: str, mime: str) -> str:
    """Pick the cleanest MIME for audio. Prefers supplied MIME, else maps
    from extension. Defaults to audio/mp3 as the most common bet."""
    if mime and mime.lower().startswith("audio/"):
        return mime.lower()
    ext = os.path.splitext(filename or "")[1].lower()
    return AUDIO_MIME_MAP.get(ext, "audio/mp3")

TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".rst", ".json", ".yaml", ".yml", ".xml",
    ".csv", ".tsv", ".log", ".html", ".htm", ".sql", ".sh", ".bat", ".ps1",
    ".ini", ".cfg", ".conf", ".toml", ".env",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".cs", ".cpp", ".c", ".h",
    ".hpp", ".go", ".rs", ".rb", ".php", ".kt", ".swift", ".scala", ".lua",
    ".dart", ".vue", ".svelte", ".css", ".scss", ".less",
}

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg"}

# Audio formats supported by Gemini 2.0 Flash audio understanding.
# Source: https://ai.google.dev/gemini-api/docs/audio
AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".webm", ".aac", ".aiff"}

AUDIO_MIME_MAP = {
    ".mp3":  "audio/mp3",
    ".wav":  "audio/wav",
    ".m4a":  "audio/m4a",
    ".ogg":  "audio/ogg",
    ".flac": "audio/flac",
    ".webm": "audio/webm",
    ".aac":  "audio/aac",
    ".aiff": "audio/aiff",
}

# OutSystems-specific extensions we recognise even if we can't fully parse them.
OUTSYSTEMS_BINARY_EXTS = {".oml", ".xif", ".eap"}


def _truncate(text: str, limit: int = MAX_TEXT_BYTES) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + f"\n\n... [truncated — original was {len(text):,} chars]"


def _decode_text(payload: bytes, meta: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
    try:
        text = payload.decode("utf-8", errors="replace")
    except Exception as e:
        meta["error"] = f"text decode failed: {e}"
        return ("", meta)
    meta["extracted"] = "text"
    return (_truncate(text), meta)


def _extract_pdf(payload: bytes, meta: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
    if pypdf is None:
        meta["error"] = "PDF support requires `pypdf` — add it to requirements.txt and reinstall."
        return ("", meta)
    try:
        reader = pypdf.PdfReader(io.BytesIO(payload))
        pages = []
        for i, page in enumerate(reader.pages):
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pages.append(f"[page {i+1}: extraction failed]")
        text = "\n\n".join(pages)
        meta["pages"] = len(reader.pages)
        meta["extracted"] = "pdf"
        return (_truncate(text), meta)
    except Exception as e:
        meta["error"] = f"pdf parse failed: {e}"
        return ("", meta)


def _extract_docx(payload: bytes, meta: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
    if docx_lib is None:
        meta["error"] = "DOCX support requires `python-docx` — add it to requirements.txt and reinstall."
        return ("", meta)
    try:
        doc = docx_lib.Document(io.BytesIO(payload))
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        # Tables as TSV-ish blocks
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                paragraphs.append(" | ".join(cells))
        text = "\n".join(paragraphs)
        meta["paragraphs"] = len(paragraphs)
        meta["extracted"] = "docx"
        return (_truncate(text), meta)
    except Exception as e:
        meta["error"] = f"docx parse failed: {e}"
        return ("", meta)


def _extract_zip(payload: bytes, meta: Dict[str, Any], *, label: str) -> Tuple[str, Dict[str, Any]]:
    """
    Crack open a zip-shaped archive (.zip, .oap). For OAP packages, we look for
    Manifest.xml first and lift its contents into the prompt. We then list every
    file (truncated by name + size), and concatenate the text-shaped contents.
    """
    try:
        zf = zipfile.ZipFile(io.BytesIO(payload))
    except zipfile.BadZipFile as e:
        meta["error"] = f"{label} archive could not be opened: {e}"
        return ("", meta)

    names = zf.namelist()
    meta["archive_kind"] = label
    meta["entry_count"] = len(names)
    meta["extracted"] = "zip"

    out_parts: List[str] = []
    out_parts.append(f"=== {label} ARCHIVE LISTING ({len(names)} entries) ===")

    # Highlight OutSystems-specific structure first (manifests, modules).
    manifest_names = [n for n in names if os.path.basename(n).lower() in (
        "manifest.xml", "applicationmanifest.xml", "package.xml", "package.json",
    )]
    oml_names = [n for n in names if n.lower().endswith(tuple(OUTSYSTEMS_BINARY_EXTS))]

    if oml_names:
        meta["outsystems_modules"] = oml_names
        out_parts.append("")
        out_parts.append(f"OutSystems modules detected ({len(oml_names)}):")
        for n in oml_names[:30]:
            try:
                size = zf.getinfo(n).file_size
            except KeyError:
                size = 0
            out_parts.append(f"  - {n}  ({size:,} bytes, binary OML — name + size only)")
        if len(oml_names) > 30:
            out_parts.append(f"  ... and {len(oml_names) - 30} more")

    if manifest_names:
        out_parts.append("")
        out_parts.append("--- Manifest content ---")
        for mn in manifest_names[:3]:
            try:
                with zf.open(mn) as fh:
                    raw = fh.read(MAX_ZIP_BYTES_PER_ENTRY)
                out_parts.append(f"\n[{mn}]\n{raw.decode('utf-8', errors='replace')}")
            except Exception as e:
                out_parts.append(f"\n[{mn}]  (could not read: {e})")

    # General-purpose listing of remaining entries.
    out_parts.append("")
    out_parts.append("--- Full entry list ---")
    for n in names[:MAX_ZIP_FILES]:
        try:
            size = zf.getinfo(n).file_size
        except KeyError:
            size = 0
        out_parts.append(f"  {n}  ({size:,} bytes)")
    if len(names) > MAX_ZIP_FILES:
        out_parts.append(f"  ... and {len(names) - MAX_ZIP_FILES} more (truncated)")

    # Concatenate text-shaped contents (excluding ones we already showed).
    shown_text_names = {n.lower() for n in manifest_names}
    text_entries = [n for n in names
                    if any(n.lower().endswith(ext) for ext in TEXT_EXTS)
                    and n.lower() not in shown_text_names]
    if text_entries:
        out_parts.append("")
        out_parts.append(f"--- Text contents ({min(len(text_entries), MAX_ZIP_FILES)} of {len(text_entries)} files) ---")
        for n in text_entries[:MAX_ZIP_FILES]:
            try:
                with zf.open(n) as fh:
                    raw = fh.read(MAX_ZIP_BYTES_PER_ENTRY)
                body = raw.decode("utf-8", errors="replace")
                out_parts.append(f"\n====FILE: {n}====")
                out_parts.append(body)
            except Exception as e:
                out_parts.append(f"\n[{n}]  (read failed: {e})")

    text = "\n".join(out_parts)
    return (_truncate(text), meta)


def extract_text(filename: str, mime: str, payload: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Public entry point. Returns (extracted_text, metadata_dict).

    metadata fields populated:
      - original_size: bytes received
      - extracted: 'text' | 'pdf' | 'docx' | 'zip' | 'image' | 'unknown'
      - error: if extraction failed
      - archive_kind, entry_count, outsystems_modules: for zips/oaps
      - pages: for pdfs
      - paragraphs: for docx
    """
    name = (filename or "").lower()
    mime_l = (mime or "").lower()
    meta: Dict[str, Any] = {"original_size": len(payload)}

    # Order matters — most specific dispatch first.
    if name.endswith(".pdf") or mime_l == "application/pdf":
        return _extract_pdf(payload, meta)
    if name.endswith(".docx") or "wordprocessingml" in mime_l:
        return _extract_docx(payload, meta)
    if name.endswith(".oap"):
        return _extract_zip(payload, meta, label="OAP")
    if name.endswith(".zip") or mime_l in ("application/zip", "application/x-zip-compressed"):
        return _extract_zip(payload, meta, label="ZIP")
    # SVGs are XML — let the text path handle them rather than vision.
    if name.endswith(".svg") or mime_l == "image/svg+xml":
        return _decode_text(payload, meta)
    if any(name.endswith(ext) for ext in IMAGE_EXTS) or mime_l.startswith("image/"):
        meta["extracted"] = "image"
        meta["note"] = "Image attached — call extract_text_async() with a Gemini API key for a vision summary."
        return ("", meta)
    if any(name.endswith(ext) for ext in AUDIO_EXTS) or mime_l.startswith("audio/"):
        meta["extracted"] = "audio"
        meta["note"] = "Audio attached — call extract_text_async() with a Gemini API key for a transcript + summary."
        return ("", meta)
    if any(name.endswith(ext) for ext in TEXT_EXTS) or mime_l.startswith("text/"):
        return _decode_text(payload, meta)

    # Fallback: try to decode as UTF-8 — many "unknown" files are actually plain text.
    try:
        sample = payload[:2048].decode("utf-8")
        meta["note"] = f"Unknown extension/MIME ({mime}); decoded as UTF-8 best-effort."
        return _decode_text(payload, meta)
    except UnicodeDecodeError:
        meta["extracted"] = "unknown"
        meta["note"] = f"Binary file ({mime or 'unknown MIME'}) — metadata only."
        return ("", meta)


async def extract_text_async(
    filename: str,
    mime: str,
    payload: bytes,
    *,
    gemini_api_key: str = "",
) -> Tuple[str, Dict[str, Any]]:
    """Async wrapper around extract_text() that adds vision extraction for images.

    For non-image inputs this just delegates to extract_text() (sync). For
    images, it calls extract_image_with_vision() when a Gemini key is provided
    — otherwise the metadata-only fallback from extract_text() is returned.

    Use this from async callers (FastAPI endpoints, agent-engine helpers)
    that have access to the Gemini key. Sync callers can keep using
    extract_text() and accept the metadata-only image fallback.
    """
    name = (filename or "").lower()
    mime_l = (mime or "").lower()

    # Detect images and audio here so we can route directly without
    # round-tripping through extract_text() for these binary cases. These
    # are the only two async branches in the extractor.
    is_image = (
        any(name.endswith(ext) for ext in IMAGE_EXTS) and not name.endswith(".svg")
    ) or (
        mime_l.startswith("image/") and mime_l != "image/svg+xml"
    )
    is_audio = (
        any(name.endswith(ext) for ext in AUDIO_EXTS)
    ) or mime_l.startswith("audio/")

    if is_image and gemini_api_key:
        meta: Dict[str, Any] = {"original_size": len(payload), "extracted": "image"}
        image_mime = _guess_image_mime(filename, mime)
        text, vmeta = await extract_image_with_vision(
            payload, image_mime, gemini_api_key=gemini_api_key, filename=filename,
        )
        meta.update(vmeta)
        if text:
            meta["extracted"] = "image_vision"
            return (_truncate(text), meta)
        meta["note"] = "Image attached; vision extraction failed — see metadata.error."
        return ("", meta)

    if is_audio and gemini_api_key:
        meta = {"original_size": len(payload), "extracted": "audio"}
        audio_mime = _guess_audio_mime(filename, mime)
        text, ameta = await extract_audio_with_gemini(
            payload, audio_mime, gemini_api_key=gemini_api_key, filename=filename,
        )
        meta.update(ameta)
        if text:
            meta["extracted"] = "audio_transcription"
            return (_truncate(text), meta)
        meta["note"] = "Audio attached; transcription failed — see metadata.error."
        return ("", meta)

    # All other paths are synchronous.
    return extract_text(filename, mime, payload)


def materials_to_prompt_block(materials: List[Dict[str, Any]],
                               *, max_total_chars: int = 80_000) -> str:
    """
    Render a list of project_materials rows (as returned by Supabase) into the
    `## Project Materials` block injected into agent prompts. Truncates the
    aggregate to max_total_chars to protect prompt budget.
    """
    if not materials:
        return ""
    lines = ["", "## Project Materials",
             "The following materials were attached to this project. Treat them as authoritative context — they are documents the team has explicitly handed you to reason against.",
             ""]
    used = sum(len(line) + 1 for line in lines)
    for m in materials:
        kind = m.get("kind", "unknown")
        title = m.get("filename") or m.get("source_url") or "(untitled)"
        body = m.get("content_text") or ""
        meta = m.get("metadata") or {}
        header_parts = [f"### [{kind}] {title}"]
        if meta.get("archive_kind"):
            header_parts.append(f" ({meta['entry_count']} entries inside the {meta['archive_kind']})")
        if meta.get("pages"):
            header_parts.append(f" ({meta['pages']} pages)")
        if meta.get("error"):
            header_parts.append(f"  ⚠ extraction error: {meta['error']}")
        header = "".join(header_parts)
        # Compute how much body we can afford.
        remaining = max_total_chars - used - len(header) - 4
        if remaining <= 200:
            lines.append(header)
            lines.append("(remaining materials trimmed to protect prompt budget)")
            break
        clipped = body if len(body) <= remaining else body[:remaining] + "\n... [trimmed]"
        lines.append(header)
        if m.get("source_url"):
            lines.append(f"Source URL: {m['source_url']}")
        if clipped:
            lines.append(clipped)
        lines.append("")
        used += len(header) + len(clipped) + 8
        if used >= max_total_chars:
            break
    return "\n".join(lines)
